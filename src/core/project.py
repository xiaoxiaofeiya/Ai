#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
项目管理核心逻辑
处理项目的创建、更新、删除等业务逻辑
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

from database.operations import DatabaseManager
from utils.logger import logger

class ProjectManager:
    """项目管理类"""
    
    def __init__(self):
        """初始化项目管理器"""
        self.db = DatabaseManager()
    
    def create_project(self, name: str, description: Optional[str] = None) -> Optional[Dict[str, Any]]:
        """创建新项目
        
        Args:
            name: 项目名称
            description: 项目描述
            
        Returns:
            项目信息字典或None（如果创建失败）
        """
        # 验证项目名称
        if not name or len(name.strip()) == 0:
            logger.error("项目名称不能为空")
            return None
        
        # 创建项目
        project = self.db.create_project(name.strip(), description)
        if not project:
            return None
        
        return {
            'id': project.id,
            'name': project.name,
            'description': project.description,
            'created_at': project.created_at,
            'updated_at': project.updated_at
        }
    
    def get_project(self, project_id: int) -> Optional[Dict[str, Any]]:
        """获取项目信息
        
        Args:
            project_id: 项目ID
            
        Returns:
            项目信息字典或None（如果项目不存在）
        """
        project = self.db.get_project(project_id)
        if not project:
            return None
            
        return {
            'id': project.id,
            'name': project.name,
            'description': project.description,
            'created_at': project.created_at,
            'updated_at': project.updated_at,
            'chapters': [
                {
                    'id': chapter.id,
                    'title': chapter.title,
                    'order': chapter.order
                }
                for chapter in project.chapters
            ]
        }
    
    def get_all_projects(self) -> List[Dict[str, Any]]:
        """获取所有项目
        
        Returns:
            项目信息列表
        """
        projects = self.db.get_all_projects()
        return [
            {
                'id': project.id,
                'name': project.name,
                'description': project.description,
                'created_at': project.created_at,
                'updated_at': project.updated_at,
                'chapter_count': len(project.chapters)
            }
            for project in projects
        ]
    
    def update_project(self, project_id: int, name: Optional[str] = None, 
                      description: Optional[str] = None) -> bool:
        """更新项目信息
        
        Args:
            project_id: 项目ID
            name: 新的项目名称
            description: 新的项目描述
            
        Returns:
            更新是否成功
        """
        update_data = {}
        if name is not None and len(name.strip()) > 0:
            update_data['name'] = name.strip()
        if description is not None:
            update_data['description'] = description
        
        if not update_data:
            logger.warning("没有需要更新的项目信息")
            return False
        
        return self.db.update_project(project_id, **update_data)
    
    def delete_project(self, project_id: int) -> bool:
        """删除项目
        
        Args:
            project_id: 项目ID
            
        Returns:
            删除是否成功
        """
        return self.db.delete_project(project_id)
    
    def export_project(self, project_id: int, export_format: str = 'json') -> Optional[str]:
        """导出项目
        
        Args:
            project_id: 项目ID
            export_format: 导出格式（json/txt/docx）
            
        Returns:
            导出文件路径或None（如果导出失败）
        """
        project = self.db.get_project(project_id)
        if not project:
            logger.error(f"项目不存在: {project_id}")
            return None
        
        # 创建导出目录
        export_dir = Path(__file__).parent.parent.parent / "exports"
        export_dir.mkdir(exist_ok=True)
        
        # 生成导出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{project.name}_{timestamp}"
        
        try:
            if export_format == 'json':
                return self._export_as_json(project, export_dir / f"{filename}.json")
            elif export_format == 'txt':
                return self._export_as_txt(project, export_dir / f"{filename}.txt")
            elif export_format == 'docx':
                return self._export_as_docx(project, export_dir / f"{filename}.docx")
            else:
                logger.error(f"不支持的导出格式: {export_format}")
                return None
        except Exception as e:
            logger.error(f"导出项目失败: {e}")
            return None
    
    def _export_as_json(self, project: Any, filepath: Path) -> str:
        """导出为JSON格式"""
        import json
        
        data = {
            'id': project.id,
            'name': project.name,
            'description': project.description,
            'created_at': project.created_at.isoformat(),
            'updated_at': project.updated_at.isoformat(),
            'chapters': [
                {
                    'id': chapter.id,
                    'title': chapter.title,
                    'content': chapter.content,
                    'order': chapter.order,
                    'created_at': chapter.created_at.isoformat(),
                    'updated_at': chapter.updated_at.isoformat()
                }
                for chapter in sorted(project.chapters, key=lambda x: x.order)
            ]
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return str(filepath)
    
    def _export_as_txt(self, project: Any, filepath: Path) -> str:
        """导出为TXT格式"""
        with open(filepath, 'w', encoding='utf-8') as f:
            # 写入项目信息
            f.write(f"# {project.name}\n\n")
            if project.description:
                f.write(f"{project.description}\n\n")
            
            # 写入章节内容
            for chapter in sorted(project.chapters, key=lambda x: x.order):
                f.write(f"## {chapter.title}\n\n")
                if chapter.content:
                    f.write(f"{chapter.content}\n\n")
        
        return str(filepath)
    
    def _export_as_docx(self, project: Any, filepath: Path) -> str:
        """导出为DOCX格式"""
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # 设置标题
        title = doc.add_heading(project.name, level=0)
        title.alignment = 1  # 居中对齐
        
        # 添加项目描述
        if project.description:
            doc.add_paragraph(project.description)
        
        # 添加章节内容
        for chapter in sorted(project.chapters, key=lambda x: x.order):
            # 添加章节标题
            doc.add_heading(chapter.title, level=1)
            
            # 添加章节内容
            if chapter.content:
                doc.add_paragraph(chapter.content)
        
        # 保存文档
        doc.save(filepath)
        
        return str(filepath)
